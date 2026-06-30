"""
Превращает markdown-текст от Клода в:
1) чистый текст без "мусорных" символов — для превью в чате (strip_markdown)
2) аккуратный Word-документ с реальными заголовками и жирным текстом (markdown_to_docx)
"""

import re

from docx import Document


def strip_markdown(text: str) -> str:
    """Убирает markdown-разметку, оставляя читаемый текст (для сообщений в Telegram,
    где жирный/заголовки всё равно не отрендерятся без parse_mode)."""
    text = re.sub(r"^#{1,6}\s+", "", text, flags=re.MULTILINE)
    text = re.sub(r"\*\*(.+?)\*\*", r"\1", text)
    text = re.sub(r"(?<!\*)\*([^*\n]+?)\*(?!\*)", r"\1", text)
    text = re.sub(r"^-{3,}\s*$", "", text, flags=re.MULTILINE)
    text = re.sub(r"^[-*]\s+", "• ", text, flags=re.MULTILINE)
    text = re.sub(r"\n{3,}", "\n\n", text)
    return text.strip()


def _strip_bold_markers(text: str) -> str:
    return re.sub(r"\*\*(.+?)\*\*", r"\1", text)


def _add_runs(paragraph, text: str) -> None:
    """Разбивает строку на обычные и **жирные** куски и добавляет их как runs."""
    parts = re.split(r"(\*\*.+?\*\*)", text)
    for part in parts:
        if not part:
            continue
        if part.startswith("**") and part.endswith("**") and len(part) > 4:
            run = paragraph.add_run(part[2:-2])
            run.bold = True
        else:
            paragraph.add_run(part)


_HEADING_RE = re.compile(
    r"^\s*(#{1,4}\s+\S.*|ЧАСТЬ\s+\S.*|\**\d+[\.\)]\s+\S.*)$",
    re.MULTILINE | re.IGNORECASE,
)


def _trim_to_sentence(text: str, limit: int) -> str:
    text = text.strip()
    if len(text) <= limit:
        return text
    trimmed = text[:limit]
    cut = trimmed.rfind(". ")
    if cut > limit * 0.4:
        return trimmed[: cut + 1]
    return trimmed


def extract_main_theme(text: str, fallback_len: int = 800) -> str:
    """Вытаскивает текст раздела "Главная тема года" из полного ответа Клода
    (заголовки у Клода каждый раз оформлены по-разному — от "1. Главная тема года"
    до "# ЧАСТЬ I. ГЛАВНАЯ ТЕМА ГОДА" — поэтому ищем гибко). Если раздел не нашёлся,
    просто берём начало текста."""
    lower = text.lower()
    idx = lower.find("главная тема года")
    if idx == -1:
        return _trim_to_sentence(strip_markdown(text), fallback_len)

    line_end = text.find("\n", idx)
    if line_end == -1:
        line_end = len(text)
    section_start = line_end + 1

    next_match = _HEADING_RE.search(text, section_start)
    section_end = next_match.start() if next_match else min(len(text), section_start + 1500)

    section_text = text[section_start:section_end]
    return _trim_to_sentence(strip_markdown(section_text), fallback_len + 400)


def markdown_to_docx(title: str, markdown_text: str, output_path: str) -> None:
    """Конвертирует markdown-текст в Word-документ с настоящими заголовками и жирным."""
    doc = Document()
    doc.add_heading(title, level=0)

    for raw_line in markdown_text.split("\n"):
        line = raw_line.rstrip()

        if not line.strip():
            continue

        if re.match(r"^-{3,}\s*$", line.strip()):
            continue

        header_match = re.match(r"^(#{1,4})\s+(.*)", line)
        if header_match:
            level = min(len(header_match.group(1)), 4)
            doc.add_heading(_strip_bold_markers(header_match.group(2)), level=level)
            continue

        bullet_match = re.match(r"^[-*]\s+(.*)", line)
        if bullet_match:
            p = doc.add_paragraph(style="List Bullet")
            _add_runs(p, bullet_match.group(1))
            continue

        numbered_match = re.match(r"^\d+\.\s+(.*)", line)
        if numbered_match:
            p = doc.add_paragraph(style="List Number")
            _add_runs(p, numbered_match.group(1))
            continue

        p = doc.add_paragraph()
        _add_runs(p, line)

    doc.save(output_path)
